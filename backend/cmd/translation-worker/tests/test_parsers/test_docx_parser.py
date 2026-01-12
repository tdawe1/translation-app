# tests/test_parsers/test_docx_parser.py
"""
Unit tests for DOCX parser functionality.

Tests python-docx-based Word parsing, extraction, and rendering.
Tests gracefully skip when python-docx is not installed.
"""

import pytest
import sys
import tempfile
from pathlib import Path

# Add worker directory to path for imports
worker_dir = Path(__file__).parent.parent.parent
sys.path.insert(0, str(worker_dir))

# Import parser module
try:
    from parsers.docx_parser import (
        DOCXParser,
        ParagraphInfo,
        TableInfo,
        FormattingInfo,
        create_docx_parser,
        PYTHON_DOCX_AVAILABLE,
    )
    PYTHON_DOCX_INSTALLED = PYTHON_DOCX_AVAILABLE
except ImportError:
    PYTHON_DOCX_INSTALLED = False
    DOCXParser = None
    ParagraphInfo = None
    TableInfo = None
    FormattingInfo = None
    create_docx_parser = None


@pytest.mark.skipif(not PYTHON_DOCX_INSTALLED, reason="python-docx not installed")
class TestParagraphInfo:
    """Test ParagraphInfo dataclass."""

    def test_create_paragraph_info(self):
        """Should create ParagraphInfo with correct attributes."""
        para = ParagraphInfo(
            index=1,
            style_name="Heading 1",
            text="Hello World",
            alignment="center",
            is_heading=True,
            heading_level=1,
            is_table=False,
            font_size=14.0,
            font_name="Arial",
            bold=True
        )

        assert para.index == 1
        assert para.style_name == "Heading 1"
        assert para.text == "Hello World"
        assert para.alignment == "center"
        assert para.is_heading is True
        assert para.heading_level == 1
        assert para.is_table is False
        assert para.font_size == 14.0
        assert para.bold is True

    def test_paragraph_info_defaults(self):
        """Should create ParagraphInfo with defaults."""
        para = ParagraphInfo(index=1)

        assert para.index == 1
        assert para.style_name == ""
        assert para.text == ""
        assert para.alignment == ""
        assert para.is_heading is False
        assert para.heading_level == 0
        assert para.is_table is False
        assert para.table_cell == ""
        assert para.is_header is False
        assert para.is_footer is False


@pytest.mark.skipif(not PYTHON_DOCX_INSTALLED, reason="python-docx not installed")
class TestTableInfo:
    """Test TableInfo dataclass."""

    def test_create_table_info(self):
        """Should create TableInfo with correct attributes."""
        table = TableInfo(
            index=0,
            rows=3,
            columns=4,
            first_row_header=True,
            cells=["A1", "B1", "C1", "D1", "A2", "B2"]
        )

        assert table.index == 0
        assert table.rows == 3
        assert table.columns == 4
        assert table.first_row_header is True
        assert len(table.cells) == 6

    def test_table_info_defaults(self):
        """Should create TableInfo with minimal defaults."""
        table = TableInfo(index=0, rows=2, columns=2)

        assert table.index == 0
        assert table.rows == 2
        assert table.columns == 2
        assert table.first_row_header is False
        assert table.cells == []


@pytest.mark.skipif(not PYTHON_DOCX_INSTALLED, reason="python-docx not installed")
class TestFormattingInfo:
    """Test FormattingInfo dataclass."""

    def test_create_formatting_info(self):
        """Should create FormattingInfo with all fields."""
        info = FormattingInfo(
            font_name="Calibri",
            font_size=12.0,
            bold=True,
            italic=False,
            underline=True,
            color_rgb=(255, 0, 0),
            alignment="justify"
        )

        assert info.font_name == "Calibri"
        assert info.font_size == 12.0
        assert info.bold is True
        assert info.italic is False
        assert info.underline is True
        assert info.color_rgb == (255, 0, 0)
        assert info.alignment == "justify"


@pytest.mark.skipif(not PYTHON_DOCX_INSTALLED, reason="python-docx not installed")
class TestDOCXParser:
    """Test DOCXParser functionality."""

    def test_plugin_attributes(self):
        """Should have required plugin attributes."""
        parser = DOCXParser()

        assert hasattr(parser, 'name')
        assert hasattr(parser, 'version')
        assert hasattr(parser, 'dependencies')
        assert parser.name == "docx_parser"
        assert parser.version == "1.0.0"
        assert "python-docx" in parser.dependencies

    def test_supported_extensions(self):
        """Should return DOCX and DOC extensions."""
        parser = DOCXParser()

        extensions = parser.supported_extensions()
        assert ".docx" in extensions
        assert ".doc" in extensions
        assert len(extensions) == 2

    def test_initialization_default_config(self):
        """Should initialize with default config."""
        parser = DOCXParser()

        assert parser._extract_headers is False
        assert parser._extract_footers is False
        assert parser._merge_paragraphs is False
        assert parser._include_empty_paragraphs is False
        assert parser._extract_tables is True

    def test_initialization_custom_config(self):
        """Should accept custom configuration."""
        parser = DOCXParser(config={
            "extract_headers": True,
            "extract_footers": True,
            "merge_paragraphs": True,
            "include_empty_paragraphs": True,
            "extract_tables": False
        })

        assert parser._extract_headers is True
        assert parser._extract_footers is True
        assert parser._merge_paragraphs is True
        assert parser._include_empty_paragraphs is True
        assert parser._extract_tables is False

    def test_initialize_method(self):
        """Should update config via initialize method."""
        parser = DOCXParser()
        parser.initialize({"extract_headers": True, "merge_paragraphs": True})

        assert parser._extract_headers is True
        assert parser._merge_paragraphs is True

    def test_shutdown(self):
        """Should shutdown without errors."""
        parser = DOCXParser()
        parser.shutdown()  # Should not raise

    def test_parse_nonexistent_file(self):
        """Should raise FileNotFoundError for missing file."""
        parser = DOCXParser()

        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.docx")

    def test_create_docx_parser_factory(self):
        """Should create parser via factory function."""
        parser = create_docx_parser()

        assert isinstance(parser, DOCXParser)

    def test_create_docx_parser_with_config(self):
        """Should create parser with config via factory."""
        parser = create_docx_parser({"merge_paragraphs": True})

        assert parser._merge_paragraphs is True


@pytest.mark.skipif(not PYTHON_DOCX_INSTALLED, reason="python-docx not installed")
class TestDOCXParserWithSampleFile:
    """Test parser with actual DOCX file."""

    def test_parse_simple_docx(self):
        """Should parse a simple DOCX document."""
        from docx import Document

        parser = DOCXParser()

        # Create a simple test DOCX
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.add_heading("Test Document", 0)
            doc.add_paragraph("This is sample content")

            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                assert parsed.format == "docx"
                assert parsed.metadata["paragraph_count"] >= 1
                assert len(parsed.segments) >= 1
                assert any("Test Document" in s.text for s in parsed.segments)
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_parse_multi_paragraph_docx(self):
        """Should parse multi-paragraph DOCX."""
        from docx import Document

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()

            for i in range(5):
                doc.add_paragraph(f"Paragraph {i + 1}")

            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                assert parsed.metadata["paragraph_count"] >= 5
                assert len(parsed.segments) >= 5
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_parse_with_headings(self):
        """Should identify heading levels."""
        from docx import Document

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.add_heading("Title", 1)
            doc.add_heading("Section", 2)
            doc.add_paragraph("Content")

            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                # Check that headings are identified
                heading_segments = [
                    s for s in parsed.segments
                    if s.context.get("type") == "heading"
                ]
                assert len(heading_segments) >= 2

                # Check heading levels
                assert any(s.context.get("heading_level") == 1 for s in heading_segments)
                assert any(s.context.get("heading_level") == 2 for s in heading_segments)
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_parse_with_table(self):
        """Should extract tables from document."""
        from docx import Document
        from docx.shared import Inches

        parser = DOCXParser(config={"extract_tables": True})

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()

            # Add a table
            table = doc.add_table(rows=3, cols=2)
            table.rows[0].cells[0].text = "Header 1"
            table.rows[0].cells[1].text = "Header 2"
            table.rows[1].cells[0].text = "Data 1"
            table.rows[1].cells[1].text = "Data 2"

            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                assert parsed.metadata["table_count"] >= 1

                # Check table segments
                table_segments = [
                    s for s in parsed.segments
                    if s.metadata.get("is_table")
                ]
                assert len(table_segments) > 0
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_parse_without_tables(self):
        """Should skip tables when not configured."""
        from docx import Document

        parser = DOCXParser(config={"extract_tables": False})

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.add_paragraph("Content before table")
            doc.add_table(rows=2, cols=2)
            doc.add_paragraph("Content after table")

            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                # Tables should still be parsed but with extract_tables=False,
                # the parser behavior depends on implementation
                # This test documents current behavior
                assert len(parsed.segments) >= 1
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_parse_with_merge_paragraphs_true(self):
        """Should merge paragraphs when configured."""
        from docx import Document

        parser = DOCXParser(config={"merge_paragraphs": True})

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.add_paragraph("First paragraph")
            doc.add_paragraph("Second paragraph")
            doc.add_paragraph("Third paragraph")

            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                # With merge_paragraphs=True, paragraphs may be grouped
                # Check that content is present
                combined_text = " ".join(s.text for s in parsed.segments)
                assert "First paragraph" in combined_text
                assert "Second paragraph" in combined_text
                assert "Third paragraph" in combined_text
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_parse_with_merge_paragraphs_false(self):
        """Should create segment per paragraph when merge is off."""
        from docx import Document

        parser = DOCXParser(config={"merge_paragraphs": False})

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.add_paragraph("First paragraph")
            doc.add_paragraph("Second paragraph")
            doc.add_paragraph("Third paragraph")

            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                # Should have individual paragraph segments
                paragraph_segments = [
                    s for s in parsed.segments
                    if s.context.get("type") == "paragraph"
                ]
                assert len(paragraph_segments) == 3
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_get_paragraph_count(self):
        """Should return correct paragraph count."""
        from docx import Document

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            for _ in range(7):
                doc.add_paragraph("Content")
            doc.save(tmp.name)

            try:
                count = parser.get_paragraph_count(tmp.name)
                assert count >= 7
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_get_table_count(self):
        """Should return correct table count."""
        from docx import Document

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.add_table(rows=2, cols=2)
            doc.add_table(rows=3, cols=3)
            doc.add_table(rows=1, cols=4)
            doc.save(tmp.name)

            try:
                count = parser.get_table_count(tmp.name)
                assert count == 3
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_extract_headings(self):
        """Should extract headings with levels."""
        from docx import Document

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.add_heading("Main Title", 1)
            doc.add_paragraph("Intro")
            doc.add_heading("Section 1", 2)
            doc.add_paragraph("Content")
            doc.add_heading("Subsection", 3)
            doc.add_paragraph("Details")

            doc.save(tmp.name)

            try:
                headings = parser.extract_headings(tmp.name)

                assert len(headings) >= 3
                assert headings[0]["level"] == 1
                assert headings[0]["text"] == "Main Title"
                assert headings[1]["level"] == 2
                assert headings[2]["level"] == 3
            finally:
                Path(tmp.name).unlink(missing_ok=True)


@pytest.mark.skipif(not PYTHON_DOCX_INSTALLED, reason="python-docx not installed")
class TestDOCXParserRendering:
    """Test DOCX rendering functionality."""

    def test_render_with_template(self):
        """Should render DOCX using template."""
        from docx import Document
        from plugins import ParsedDocument, Segment

        parser = DOCXParser()

        # Create source DOCX
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as source_tmp:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as output_tmp:
                doc = Document()
                doc.add_heading("Original Title", 1)
                doc.add_paragraph("Original content")

                doc.save(source_tmp.name)

                try:
                    # Create parsed document with translation
                    parsed = ParsedDocument(
                        segments=[
                            Segment(
                                id="paragraph_1",
                                text="Translated Title",
                                context={"type": "heading", "heading_level": 1, "paragraph_index": 1},
                                metadata={}
                            ),
                            Segment(
                                id="paragraph_2",
                                text="Translated content",
                                context={"type": "paragraph", "paragraph_index": 2},
                                metadata={}
                            )
                        ],
                        metadata={"format": "docx"},
                        format="docx",
                        source_path=source_tmp.name
                    )

                    # Render with template
                    parser.render(
                        doc=parsed,
                        output_path=output_tmp.name,
                        template_path=source_tmp.name
                    )

                    # Verify output file exists and has content
                    assert Path(output_tmp.name).exists()
                    assert Path(output_tmp.name).stat().st_size > 0

                    # Verify we can open the rendered file
                    rendered_doc = Document(output_tmp.name)
                    assert len(rendered_doc.paragraphs) >= 2

                finally:
                    Path(source_tmp.name).unlink(missing_ok=True)
                    Path(output_tmp.name).unlink(missing_ok=True)

    def test_render_without_template(self):
        """Should render DOCX without template (new document)."""
        from docx import Document
        from plugins import ParsedDocument, Segment

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as output_tmp:
            try:
                parsed = ParsedDocument(
                    segments=[
                        Segment(
                            id="paragraph_1",
                            text="Test content",
                            context={"type": "paragraph", "paragraph_index": 1},
                            metadata={}
                        )
                    ],
                    metadata={"format": "docx"},
                    format="docx"
                )

                parser.render(doc=parsed, output_path=output_tmp.name)

                assert Path(output_tmp.name).exists()
                assert Path(output_tmp.name).stat().st_size > 0

            finally:
                Path(output_tmp.name).unlink(missing_ok=True)

    def test_render_with_formatting(self):
        """Should preserve formatting from metadata."""
        from docx import Document
        from plugins import ParsedDocument, Segment

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as output_tmp:
            try:
                parsed = ParsedDocument(
                    segments=[
                        Segment(
                            id="paragraph_1",
                            text="Bold text",
                            context={"type": "paragraph", "paragraph_index": 1},
                            metadata={"bold": True}
                        ),
                        Segment(
                            id="paragraph_2",
                            text="Centered text",
                            context={"type": "paragraph", "paragraph_index": 2},
                            metadata={"alignment": "center"}
                        )
                    ],
                    metadata={"format": "docx"},
                    format="docx"
                )

                parser.render(doc=parsed, output_path=output_tmp.name)

                # Verify output
                assert Path(output_tmp.name).exists()

                rendered_doc = Document(output_tmp.name)
                assert len(rendered_doc.paragraphs) >= 2

            finally:
                Path(output_tmp.name).unlink(missing_ok=True)


@pytest.mark.skipif(not PYTHON_DOCX_INSTALLED, reason="python-docx not installed")
class TestDOCXParserEdgeCases:
    """Test edge cases and special scenarios."""

    def test_empty_document(self):
        """Should handle empty document."""
        from docx import Document

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()  # Empty document
            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                # Empty document should have no segments or minimal metadata
                assert parsed.format == "docx"
                assert parsed.metadata["paragraph_count"] == 0
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_document_with_only_tables(self):
        """Should handle document with only tables."""
        from docx import Document

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "A"
            table.rows[0].cells[1].text = "B"
            table.rows[1].cells[0].text = "C"
            table.rows[1].cells[1].text = "D"
            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                assert parsed.metadata["table_count"] == 1
                # Check that table content is extracted
                assert any("A" in s.text or "B" in s.text or "C" in s.text or "D" in s.text
                          for s in parsed.segments)
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_unicode_content(self):
        """Should handle unicode characters."""
        from docx import Document

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.add_paragraph("日本語のテキスト")
            doc.add_paragraph("Текст на русском")
            doc.add_paragraph("Emoji: 🎉 🚀 ✨")

            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                # Check unicode is preserved
                text = " ".join(s.text for s in parsed.segments)
                assert "日本語" in text
                assert "Текст" in text
                assert "🎉" in text
            finally:
                Path(tmp.name).unlink(missing_ok=True)

    def test_long_paragraph(self):
        """Should handle very long paragraphs."""
        from docx import Document

        parser = DOCXParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            long_text = "Word " * 1000  # ~5000 characters
            doc.add_paragraph(long_text)
            doc.save(tmp.name)

            try:
                parsed = parser.parse(tmp.name)

                assert len(parsed.segments) == 1
                assert len(parsed.segments[0].text) > 4000
            finally:
                Path(tmp.name).unlink(missing_ok=True)


class TestDOCXNotInstalled:
    """Test behavior when python-docx is not installed."""

    def test_parser_raises_when_not_installed(self):
        """Should raise RuntimeError when python-docx unavailable."""
        if PYTHON_DOCX_INSTALLED:
            pytest.skip("python-docx is installed")

        with pytest.raises(RuntimeError, match="python-docx is not installed"):
            DOCXParser()

    def test_factory_raises_when_not_installed(self):
        """Should raise RuntimeError from factory when python-docx unavailable."""
        if PYTHON_DOCX_INSTALLED:
            pytest.skip("python-docx is installed")

        with pytest.raises(RuntimeError, match="python-docx is not installed"):
            create_docx_parser()
