# parsers/docx_parser.py
"""
DOCX parser using python-docx for Word documents.

Extracts text from documents with formatting and structure preservation.
Handles paragraphs, tables, headers, footers, and styles.
"""

from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
import logging

logger = logging.getLogger(__name__)

# Try to import python-docx
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    Document = None

# Import from plugins module
import sys
from pathlib import Path

worker_dir = Path(__file__).parent.parent
sys.path.insert(0, str(worker_dir))

from plugins import (
    ParserPlugin,
    ParsedDocument,
    Segment,
)


@dataclass
class ParagraphInfo:
    """Information about a Word paragraph."""

    index: int  # 0-based paragraph index
    style_name: str = ""
    text: str = ""
    alignment: str = ""  # left, center, right, justify
    is_heading: bool = False
    heading_level: int = 0  # 1-9 for headings, 0 for normal text
    is_table: bool = False
    table_cell: str = ""  # "A1", "B2", etc. if in table
    is_header: bool = False
    is_footer: bool = False
    section_number: int = 0

    # Formatting
    font_size: float = 0.0
    font_name: str = ""
    bold: bool = False
    italic: bool = False
    underline: bool = False


@dataclass
class TableInfo:
    """Information about a Word table."""

    index: int
    rows: int
    columns: int
    first_row_header: bool = False
    cells: List[str] = field(default_factory=list)  # Flattened cell contents


@dataclass
class FormattingInfo:
    """Text formatting information."""

    font_name: str = ""
    font_size: float = 0.0
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color_rgb: Optional[Tuple[int, int, int]] = None
    highlight_color: Optional[str] = None
    alignment: str = ""  # left, center, right, justify


class DOCXParser:
    """Word DOCX parser using python-docx.

    Features:
    - Extract text from paragraphs with position context
    - Preserve formatting (bold, italic, font size, alignment)
    - Handle tables with cell-level granularity
    - Extract headers and footers
    - Identify heading levels and document structure
    - Track paragraph styles

    Note: python-docx must be installed. Install with: pip install python-docx
    """

    name = "docx_parser"
    version = "1.0.0"
    dependencies = ["python-docx"]

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """Initialize DOCX parser.

        Args:
            config: Optional configuration dict

        Raises:
            RuntimeError: If python-docx is not installed
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError(
                "python-docx is not installed. Install with: pip install python-docx"
            )

        self.config = config or {}
        self._extract_headers = self.config.get("extract_headers", False)
        self._extract_footers = self.config.get("extract_footers", False)
        self._merge_paragraphs = self.config.get("merge_paragraphs", False)
        self._include_empty_paragraphs = self.config.get(
            "include_empty_paragraphs", False
        )
        self._extract_tables = self.config.get("extract_tables", True)

    def initialize(self, config: Dict[str, Any]) -> None:
        """Initialize plugin with configuration.

        Args:
            config: Configuration dict
        """
        self.config.update(config)
        self._extract_headers = self.config.get("extract_headers", False)
        self._extract_footers = self.config.get("extract_footers", False)
        self._merge_paragraphs = self.config.get("merge_paragraphs", False)
        self._include_empty_paragraphs = self.config.get(
            "include_empty_paragraphs", False
        )
        self._extract_tables = self.config.get("extract_tables", True)
        logger.info(f"[{self.name}] Initialized with config: {self.config}")

    def shutdown(self) -> None:
        """Clean up resources."""
        logger.debug(f"[{self.name}] Shutdown")

    def supported_extensions(self) -> List[str]:
        """Return list of supported file extensions.

        Returns:
            List of supported extensions
        """
        return [".docx", ".doc"]

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse DOCX and extract translatable text segments.

        Args:
            file_path: Path to DOCX file

        Returns:
            ParsedDocument with segments

        Raises:
            FileNotFoundError: If file doesn't exist
            Exception: If DOCX cannot be parsed
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc = Document(file_path)

        try:
            segments = []
            metadata = {
                "format": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "title": self._get_document_title(doc),
            }

            # Extract paragraphs
            paragraphs = self._extract_paragraphs_info(doc)

            # Extract tables
            tables = []
            if self._extract_tables:
                tables = self._extract_tables_info(doc)
                metadata["table_details"] = [
                    {"rows": t.rows, "columns": t.columns} for t in tables
                ]

            # Create segments from paragraphs
            paragraph_segments = self._create_paragraph_segments(paragraphs, tables)
            segments.extend(paragraph_segments)

            # Extract headers if configured
            if self._extract_headers:
                header_segments = self._extract_headers_segments(doc)
                segments.extend(header_segments)
                metadata["header_count"] = len(header_segments)

            # Extract footers if configured
            if self._extract_footers:
                footer_segments = self._extract_footers_segments(doc)
                segments.extend(footer_segments)
                metadata["footer_count"] = len(footer_segments)

            return ParsedDocument(
                segments=segments,
                metadata=metadata,
                format="docx",
                source_path=file_path,
            )

        except Exception as e:
            logger.error(f"[{self.name}] Failed to parse DOCX: {e}")
            raise
        finally:
            # Document auto-closes
            pass

    def _get_document_title(self, doc) -> str:
        """Extract document title from core properties or first paragraph.

        Args:
            doc: python-docx Document object

        Returns:
            Document title
        """
        # Try core properties first
        if hasattr(doc.core_properties, "title") and doc.core_properties.title:
            return doc.core_properties.title

        # Fall back to first non-empty paragraph
        for para in doc.paragraphs[:3]:  # Check first 3 paragraphs
            text = para.text.strip()
            if text and len(text) < 100:
                return text

        return ""

    def _extract_paragraphs_info(self, doc) -> List[ParagraphInfo]:
        """Extract information about all paragraphs.

        Args:
            doc: python-docx Document object

        Returns:
            List of ParagraphInfo objects
        """
        paragraphs = []
        idx = 0

        for para in doc.paragraphs:
            if not para.text.strip() and not self._include_empty_paragraphs:
                continue

            is_heading, heading_level = self._get_heading_info(para)
            alignment = self._get_alignment_name(para.alignment)
            formatting = self._extract_formatting(para)

            paragraphs.append(
                ParagraphInfo(
                    index=idx,
                    style_name=para.style.name if para.style else "",
                    text=para.text,
                    alignment=alignment,
                    is_heading=is_heading,
                    heading_level=heading_level,
                    is_table=False,
                    table_cell="",
                    font_size=formatting.font_size,
                    font_name=formatting.font_name,
                    bold=formatting.bold,
                    italic=formatting.italic,
                    underline=formatting.underline,
                )
            )
            idx += 1

        if self._extract_tables:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            if (
                                not para.text.strip()
                                and not self._include_empty_paragraphs
                            ):
                                continue

                            cell_ref = f"{chr(65 + col_idx)}{row_idx + 1}"
                            formatting = self._extract_formatting(para)

                            paragraphs.append(
                                ParagraphInfo(
                                    index=idx,
                                    style_name=para.style.name if para.style else "",
                                    text=para.text,
                                    alignment=self._get_alignment_name(para.alignment),
                                    is_heading=False,
                                    heading_level=0,
                                    is_table=True,
                                    table_cell=cell_ref,
                                    font_size=formatting.font_size,
                                    font_name=formatting.font_name,
                                    bold=formatting.bold,
                                    italic=formatting.italic,
                                    underline=formatting.underline,
                                )
                            )
                            idx += 1

        return paragraphs

    def _is_paragraph_in_table(self, para, tables) -> bool:
        """Check if paragraph is inside a table.

        Args:
            para: Paragraph object
            tables: List of tables in document

        Returns:
            True if paragraph is in a table
        """
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for cell_para in cell.paragraphs:
                        if cell_para == para:
                            return True
        return False

    def _get_table_cell_reference(self, para, tables) -> str:
        """Get cell reference (e.g., "A1", "B2") for paragraph in table.

        Args:
            para: Paragraph object
            tables: List of tables in document

        Returns:
            Cell reference string or empty string
        """
        for table_idx, table in enumerate(tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for cell_para in cell.paragraphs:
                        if cell_para == para:
                            # Convert to Excel-style column letter + row number
                            col_letter = (
                                chr(65 + cell_idx)
                                if cell_idx < 26
                                else f"A{cell_idx - 26}"
                            )
                            return f"table_{table_idx + 1}_{col_letter}{row_idx + 1}"
        return ""

    def _get_heading_info(self, para) -> Tuple[bool, int]:
        """Determine if paragraph is a heading and its level.

        Args:
            para: Paragraph object

        Returns:
            Tuple of (is_heading, heading_level)
        """
        if not para.style:
            return False, 0

        style_name = para.style.name.lower()

        # Check for heading styles
        if "heading" in style_name or "title" in style_name:
            # Extract level number
            for i in range(1, 10):
                if str(i) in style_name or f"h{i}" in style_name:
                    return True, i
            return True, 1  # Default to level 1

        return False, 0

    def _get_alignment_name(self, alignment) -> str:
        """Convert alignment enum to string.

        Args:
            alignment: WD_ALIGN_PARAGRAPH enum or None

        Returns:
            Alignment name string
        """
        if alignment is None:
            return "left"

        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        return alignment_map.get(alignment, "left")

    def _extract_formatting(self, para) -> FormattingInfo:
        """Extract formatting from paragraph.

        Args:
            para: python-docx Paragraph object

        Returns:
            FormattingInfo with formatting details
        """
        info = FormattingInfo()

        if para.runs:
            # Use first run's formatting
            run = para.runs[0]
            if run.font:
                info.font_name = run.font.name
                info.bold = run.font.bold
                info.italic = run.font.italic
                info.underline = run.font.underline

                size_pt = run.font.size
                if size_pt:
                    info.font_size = size_pt.pt

                # Extract color if available
                if run.font.color and run.font.color.rgb:
                    info.color_rgb = (
                        run.font.color.rgb[0],
                        run.font.color.rgb[1],
                        run.font.color.rgb[2],
                    )

        return info

    def _extract_tables_info(self, doc) -> List[TableInfo]:
        """Extract information about all tables.

        Args:
            doc: python-docx Document object

        Returns:
            List of TableInfo objects
        """
        tables = []

        for idx, table in enumerate(doc.tables):
            rows = len(table.rows)
            columns = len(table.columns)

            # Extract cell contents
            cells = []
            for row in table.rows:
                for cell in row.cells:
                    cells.append(cell.text.strip())

            # Check if first row is likely a header
            first_row_header = False
            if rows > 1 and table.rows:
                first_row_text = [c.text.strip() for c in table.rows[0].cells]
                # If all cells in first row are relatively short, likely headers
                if all(len(t) < 50 for t in first_row_text if t):
                    first_row_header = True

            tables.append(
                TableInfo(
                    index=idx,
                    rows=rows,
                    columns=columns,
                    first_row_header=first_row_header,
                    cells=cells,
                )
            )

        return tables

    def _create_paragraph_segments(
        self, paragraphs: List[ParagraphInfo], tables: List[TableInfo]
    ) -> List[Segment]:
        """Create segments from paragraph information.

        Args:
            paragraphs: List of paragraph info
            tables: List of table info

        Returns:
            List of Segment objects
        """
        segments = []

        if self._merge_paragraphs:
            # Merge consecutive non-table paragraphs
            merged_text = []
            merged_start_idx = 0

            for idx, para in enumerate(paragraphs):
                if para.is_table:
                    # Flush any accumulated merged text
                    if merged_text:
                        text = "\n".join(merged_text)
                        if text.strip():
                            segments.append(
                                Segment(
                                    id=f"paragraph_{merged_start_idx + 1}_to_{idx}",
                                    text=text,
                                    context={
                                        "type": "paragraph_group",
                                        "start_index": merged_start_idx + 1,
                                        "end_index": idx,
                                    },
                                    metadata={},
                                )
                            )
                        merged_text = []
                        merged_start_idx = idx + 1

                    # Add table as separate segment
                    if para.text.strip():
                        segments.append(
                            Segment(
                                id=f"table_{para.table_cell}",
                                text=para.text,
                                context={
                                    "type": "table",
                                    "cell_reference": para.table_cell,
                                },
                                metadata={
                                    "is_table": True,
                                    "table_cell": para.table_cell,
                                },
                            )
                        )
                else:
                    merged_text.append(para.text)

            # Flush remaining merged text
            if merged_text:
                text = "\n".join(merged_text)
                if text.strip():
                    segments.append(
                        Segment(
                            id=f"paragraph_{merged_start_idx + 1}_to_end",
                            text=text,
                            context={
                                "type": "paragraph_group",
                                "start_index": merged_start_idx + 1,
                            },
                            metadata={},
                        )
                    )
        else:
            # Create segment per paragraph
            for para in paragraphs:
                if para.text.strip():
                    seg_type = (
                        "table"
                        if para.is_table
                        else ("heading" if para.is_heading else "paragraph")
                    )

                    segments.append(
                        Segment(
                            id=f"paragraph_{para.index + 1}",
                            text=para.text,
                            context={
                                "type": seg_type,
                                "paragraph_index": para.index + 1,
                                "style_name": para.style_name,
                                "heading_level": para.heading_level
                                if para.is_heading
                                else None,
                            },
                            metadata={
                                "alignment": para.alignment,
                                "font_size": para.font_size,
                                "font_name": para.font_name,
                                "bold": para.bold,
                                "italic": para.italic,
                                "underline": para.underline,
                                "is_table": para.is_table,
                                "table_cell": para.table_cell
                                if para.is_table
                                else None,
                            },
                        )
                    )

        return segments

    def _extract_headers_segments(self, doc) -> List[Segment]:
        """Extract header text segments.

        Args:
            doc: python-docx Document object

        Returns:
            List of header segments
        """
        segments = []

        for section_idx, section in enumerate(doc.sections):
            # Headers come in first/odd/even page variants
            for header_type, header in [
                ("first", section.first_page_header),
                ("primary", section.header),
                ("even", section.even_page_header),
            ]:
                if header and hasattr(header, "paragraphs"):
                    for para_idx, para in enumerate(header.paragraphs):
                        if para.text.strip():
                            segments.append(
                                Segment(
                                    id=f"section_{section_idx + 1}_header_{header_type}_{para_idx + 1}",
                                    text=para.text,
                                    context={
                                        "type": "header",
                                        "section": section_idx + 1,
                                        "header_type": header_type,
                                        "paragraph_index": para_idx + 1,
                                    },
                                    metadata={"is_header": True},
                                )
                            )

        return segments

    def _extract_footers_segments(self, doc) -> List[Segment]:
        """Extract footer text segments.

        Args:
            doc: python-docx Document object

        Returns:
            List of footer segments
        """
        segments = []

        for section_idx, section in enumerate(doc.sections):
            # Footers come in first/odd/even page variants
            for footer_type, footer in [
                ("first", section.first_page_footer),
                ("primary", section.footer),
                ("even", section.even_page_footer),
            ]:
                if footer and hasattr(footer, "paragraphs"):
                    for para_idx, para in enumerate(footer.paragraphs):
                        if para.text.strip():
                            segments.append(
                                Segment(
                                    id=f"section_{section_idx + 1}_footer_{footer_type}_{para_idx + 1}",
                                    text=para.text,
                                    context={
                                        "type": "footer",
                                        "section": section_idx + 1,
                                        "footer_type": footer_type,
                                        "paragraph_index": para_idx + 1,
                                    },
                                    metadata={"is_footer": True},
                                )
                            )

        return segments

    def render(
        self, doc: ParsedDocument, output_path: str, template_path: Optional[str] = None
    ) -> None:
        """Render translated DOCX with original layout.

        Args:
            doc: ParsedDocument with translated segments
            output_path: Where to save the output DOCX
            template_path: Optional original DOCX as template

        Raises:
            Exception: If rendering fails
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed")

        # Load template or create new document
        if template_path:
            try:
                out_doc = Document(template_path)
            except Exception as e:
                logger.warning(f"[{self.name}] Failed to load template: {e}")
                out_doc = Document()
        else:
            out_doc = Document()

        try:
            # Map segments by paragraph index for easy lookup
            paragraph_map: Dict[int, List[Segment]] = {}
            table_map: Dict[str, List[Segment]] = {}
            header_map: Dict[str, List[Segment]] = {}
            footer_map: Dict[str, List[Segment]] = {}

            for segment in doc.segments:
                seg_type = segment.context.get("type", "paragraph")

                if seg_type == "table":
                    cell_ref = segment.metadata.get("table_cell", "")
                    if cell_ref not in table_map:
                        table_map[cell_ref] = []
                    table_map[cell_ref].append(segment)

                elif seg_type == "header":
                    key = f"{segment.context.get('section', 0)}_{segment.context.get('header_type', 'primary')}"
                    if key not in header_map:
                        header_map[key] = []
                    header_map[key].append(segment)

                elif seg_type == "footer":
                    key = f"{segment.context.get('section', 0)}_{segment.context.get('footer_type', 'primary')}"
                    if key not in footer_map:
                        footer_map[key] = []
                    footer_map[key].append(segment)

                elif seg_type == "paragraph_group":
                    # Paragraph groups are merged - add to first paragraph
                    start_idx = segment.context.get("start_index", 1)
                    paragraph_map[start_idx] = [segment]

                else:
                    # Regular paragraph or heading
                    para_idx = segment.context.get("paragraph_index", 1)
                    if para_idx not in paragraph_map:
                        paragraph_map[para_idx] = []
                    paragraph_map[para_idx].append(segment)

            # Clear and replace paragraphs
            if template_path:
                # Replace existing paragraphs
                for para_idx, para in enumerate(out_doc.paragraphs):
                    segments = paragraph_map.get(para_idx + 1, [])
                    if segments and segments[0].text:
                        para.text = segments[0].text
                        # Apply formatting from metadata
                        self._apply_paragraph_formatting(para, segments[0])
            else:
                # Create new document from segments
                for segment in doc.segments:
                    if segment.context.get("type") in ("header", "footer"):
                        continue  # Skip headers/footers for new doc

                    para = out_doc.add_paragraph(segment.text)
                    self._apply_paragraph_formatting(para, segment)

            # Handle tables if template has them
            if template_path and table_map:
                for table_idx, table in enumerate(out_doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            cell_ref = f"table_{table_idx + 1}_{chr(65 + cell_idx) if cell_idx < 26 else 'A'}{row_idx + 1}"
                            segments = table_map.get(cell_ref, [])
                            if segments and segments[0].text:
                                cell.text = segments[0].text

            out_doc.save(output_path)
            logger.info(f"[{self.name}] Rendered DOCX to: {output_path}")

        except Exception as e:
            logger.error(f"[{self.name}] Failed to render DOCX: {e}")
            raise

    def _apply_paragraph_formatting(self, para, segment: Segment) -> None:
        """Apply formatting from segment metadata to paragraph.

        Args:
            para: python-docx Paragraph object
            segment: Segment with formatting metadata
        """
        metadata = segment.metadata

        # Apply alignment
        alignment = metadata.get("alignment", "left")
        if alignment == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "justify":
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Apply run formatting if runs exist
        if para.runs:
            run = para.runs[0]
            if metadata.get("bold"):
                run.font.bold = True
            if metadata.get("italic"):
                run.font.italic = True
            if metadata.get("underline"):
                run.font.underline = True
            if metadata.get("font_size"):
                run.font.size = metadata["font_size"]
            if metadata.get("font_name"):
                run.font.name = metadata["font_name"]

    def get_paragraph_count(self, file_path: str) -> int:
        """Get the number of paragraphs in a document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Number of paragraphs
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed")

        doc = Document(file_path)
        return len(doc.paragraphs)

    def get_table_count(self, file_path: str) -> int:
        """Get the number of tables in a document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Number of tables
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed")

        doc = Document(file_path)
        return len(doc.tables)

    def extract_headings(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract headings from document.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of heading dicts with text, level, and index
        """
        doc = Document(file_path)
        headings = []

        for idx, para in enumerate(doc.paragraphs):
            is_heading, level = self._get_heading_info(para)
            if is_heading and para.text.strip():
                headings.append(
                    {
                        "text": para.text.strip(),
                        "level": level,
                        "index": idx,
                    }
                )

        return headings


def create_docx_parser(config: Optional[Dict[str, Any]] = None) -> DOCXParser:
    """Factory function to create a DOCX parser.

    Args:
        config: Optional configuration dict

    Returns:
        DOCXParser instance

    Raises:
        RuntimeError: If python-docx is not installed
    """
    return DOCXParser(config)
